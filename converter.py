from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, size, is_bold=True):
    """
    Adds a centered heading with controlled spacing.
    Used for the top title section of the form.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.0

    run = para.add_run(text)
    run.bold = is_bold
    run.font.size = Pt(size)


def write_cell(cell, text, bold=False):
    """
    Writes text inside a table cell with consistent spacing.
    This helps maintain uniform row height across the table.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.2

    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def generate_docx_from_pdf(pdf_path):

    output_path = "outputs/Mediation_Application_Form_Replica.docx"
    # Create the document
    doc = Document()

    # Adjust page margins to better match the PDF layout
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title section
    add_heading(doc, "FORM 'A'", 13, True)
    add_heading(doc, "MEDIATION APPLICATION FORM", 13, True)
    add_heading(doc, "[REFER RULE 3(1)]", 13, True)
    add_heading(doc, "Mumbai District Legal Services Authority", 12, False)
    add_heading(doc, "City Civil Court, Mumbai", 12, False)

    doc.add_paragraph()  # small gap before table

    # Create main table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Column widths based on PDF proportions
    table.columns[0].width = Inches(0.35)
    table.columns[1].width = Inches(1.2)
    table.columns[2].width = Inches(3.3)

    # DETAILS OF PARTIES section
    row = table.rows[0].cells
    write_cell(row[0], "DETAILS OF PARTIES:", bold=True)
    row[0].merge(row[1]).merge(row[2])

    # Applicant name
    row = table.add_row().cells
    write_cell(row[0], "1")
    write_cell(row[1], "Name of Applicant", bold=True)
    write_cell(row[2], "{{client_name}}", bold=True)

    # Applicant address heading
    row = table.add_row().cells
    write_cell(row[1], "Address and contact details of Applicant", bold=True)
    row[1].merge(row[2])

    # Applicant address block
    row = table.add_row().cells
    write_cell(row[1], "Address", bold=True)

    cell = row[2]
    cell.text = ""
    para = cell.paragraphs[0]

    run = para.add_run("REGISTERED ADDRESS:\n")
    run.bold = True
    run.font.size = Pt(10)

    run = para.add_run("{{branch_address}}\n\n\n")
    run.font.size = Pt(10)

    run = para.add_run("CORRESPONDENCE BRANCH ADDRESS:\n")
    run.bold = True
    run.font.size = Pt(10)

    run = para.add_run("{{branch_address}}")
    run.font.size = Pt(10)

    # Contact details
    row = table.add_row().cells
    write_cell(row[1], "Telephone No.", bold=True)
    write_cell(row[2], "{{mobile}}", bold=True)

    row = table.add_row().cells
    write_cell(row[1], "Mobile No.", bold=True)
    write_cell(row[2], "")

    row = table.add_row().cells
    write_cell(row[1], "Email ID", bold=True)
    write_cell(row[2], "info@kslegal.co.in")

    # Opposite party section
    row = table.add_row().cells
    write_cell(row[0], "2")
    write_cell(row[1], "Name, Address and Contact details of Opposite Party:", bold=True)
    row[1].merge(row[2])

    row = table.add_row().cells
    write_cell(row[1], "Address and contact details of Defendant/s", bold=True)
    row[1].merge(row[2])

    row = table.add_row().cells
    write_cell(row[1], "Name", bold=True)
    write_cell(row[2], "{{customer_name}}", bold=True)

    # Defendant address
    row = table.add_row().cells
    write_cell(row[1], "Address", bold=True)

    cell = row[2]
    cell.text = ""
    para = cell.paragraphs[0]

    run = para.add_run("REGISTERED ADDRESS:\n")
    run.bold = True
    run.font.size = Pt(10)

    run = para.add_run("{% if address1 and address1 != \"\" %}{{address1}} {% else %}_______________{% endif %}\n\n")
    run.font.size = Pt(10)

    run = para.add_run("CORRESPONDENCE ADDRESS:\n")
    run.bold = True
    run.font.size = Pt(10)

    run = para.add_run("{% if address1 and address1 != \"\" %}{{address1}} {% else %}_______________{% endif %}")
    run.font.size = Pt(10)

    # Contact rows
    for field in ["Telephone No.", "Mobile No.", "Email ID"]:
        row = table.add_row().cells
        write_cell(row[1], field, bold=True)
        write_cell(row[2], "")

    # DETAILS OF DISPUTE section
    row = table.add_row().cells
    write_cell(row[0], "DETAILS OF DISPUTE:", bold=True)
    row[0].merge(row[1]).merge(row[2])

    row = table.add_row().cells
    cell = row[0]
    cell.text = ""

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)

    run = para.add_run("THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES, 2018")
    run.bold = True
    run.font.size = Pt(10)

    row[0].merge(row[1]).merge(row[2])

    row = table.add_row().cells
    write_cell(
        row[1],
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):",
        bold=True
    )
    row[1].merge(row[2])

# Save document
    doc.save(output_path)
    return output_path
